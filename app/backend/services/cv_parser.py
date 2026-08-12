import io
import re
from pathlib import Path
from typing import Any

import fitz
from docx import Document


SUPPORTED_EXTENSIONS = {".pdf", ".docx"}

COMMON_SKILLS = {
    "python",
    "java",
    "javascript",
    "typescript",
    "c++",
    "c#",
    "sql",
    "html",
    "css",
    "react",
    "angular",
    "vue",
    "node.js",
    "fastapi",
    "flask",
    "django",
    "streamlit",
    "postgresql",
    "mysql",
    "mongodb",
    "sqlite",
    "git",
    "github",
    "docker",
    "kubernetes",
    "aws",
    "azure",
    "gcp",
    "pandas",
    "numpy",
    "scikit-learn",
    "tensorflow",
    "pytorch",
    "machine learning",
    "deep learning",
    "natural language processing",
    "nlp",
    "large language models",
    "llm",
    "power bi",
    "tableau",
    "excel",
    "project management",
    "agile",
    "scrum",
    "communication",
    "leadership",
    "problem solving",
}

class CVParserError(Exception):
    """Raised when a CV cannot be parsed."""


def clean_text(text: str) -> str:
    """
    Clean extracted CV text while preserving useful line breaks.
    """
    if not text:
        return ""

    text = text.replace("\x00", " ")  #replace null characters with space
    text = text.replace("\r\n", "\n") #replace line separators with new line
    text = text.replace("\r", "\n") #replace remaining spaces with new line

    # Replace repeated spaces and tabs.
    text = re.sub(r"[ \t]+", " ", text)

    # Remove excessive blank lines.
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Strip whitespace from each line.
    lines = [line.strip() for line in text.splitlines()]
    lines = [line for line in lines if line]

    return "\n".join(lines).strip()


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text from a PDF file using PyMuPDF.
    """
    try:
        with fitz.open(stream=file_bytes, filetype="pdf") as document:
            pages: list[str] = []

            for page in document:
                page_text = page.get_text("text", sort=True)
                if page_text:
                    pages.append(page_text)

        return clean_text("\n".join(pages))

    except Exception as exc:
        raise CVParserError(f"Unable to read PDF file: {exc}") from exc


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text from DOCX paragraphs and tables.
    """
    try:
        document = Document(io.BytesIO(file_bytes))
        content: list[str] = []

        # Extract paragraphs.
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                content.append(text)

        # Extract table content.
        for table in document.tables:
            for row in table.rows:
                row_values = [
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                ]

                if row_values:
                    content.append(" | ".join(row_values))

        return clean_text("\n".join(content))

    except Exception as exc:
        raise CVParserError(f"Unable to read DOCX file: {exc}") from exc

def extract_email(text: str) -> str | None:
    """
    Extract the first email address found in the CV.
    """
    pattern = r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b"
    match = re.search(pattern, text)

    return match.group(0) if match else None


def extract_phone(text: str) -> str | None:
    """
    Extract a possible phone number.
    """
    pattern = r"(?<!\d)(?:\+?\d{1,3}[\s\-()]*)?(?:\d[\s\-()]*){9,12}(?!\d)"
    match = re.search(pattern, text)

    if not match:
        return None

    return re.sub(r"\s+", " ", match.group(0)).strip()


def extract_skills(text: str) -> list[str]:
    """
    Find skills from the predefined skill list.
    """
    normalised_text = text.lower()
    detected_skills: list[str] = []

    for skill in COMMON_SKILLS:
        pattern = rf"(?<!\w){re.escape(skill.lower())}(?!\w)"

        if re.search(pattern, normalised_text):
            detected_skills.append(skill)

    return sorted(detected_skills)


def extract_section(
    text: str,
    section_names: list[str],
    stopping_sections: list[str],
) -> str | None:
    """
    Extract text beneath a CV section heading.

    Example headings:
    Experience
    Work Experience
    Work History
    Education
    Skills
    Projects
    """
    lines = text.splitlines()
    start_index: int | None = None
    collected: list[str] = []

    section_names_lower = {
        name.lower().strip(":")
        for name in section_names
    }

    stopping_sections_lower = {
        name.lower().strip(":")
        for name in stopping_sections
    }

    for index, line in enumerate(lines):
        heading = line.lower().strip().strip(":")

        if heading in section_names_lower:
            start_index = index + 1
            break

    if start_index is None:
        return None

    for line in lines[start_index:]:
        heading = line.lower().strip().strip(":")

        if heading in stopping_sections_lower:
            break

        collected.append(line)

    section_text = "\n".join(collected).strip()

    return section_text or None


def extract_education(text: str) -> str | None:
    return extract_section(
        text=text,
        section_names=[
            "education",
            "academic background",
            "academic qualifications",
            "qualifications",
        ],
        stopping_sections=[
            "experience",
            "work experience",
            "employment history",
            "skills",
            "technical skills",
            "projects",
            "certifications",
            "references",
        ],
    )
def extract_experience(text: str) -> str | None:
    return extract_section(
        text=text,
        section_names=[
            "experience",
            "work experience",
            "work history",
            "professional experience",
            "employment history",
            "career history",
        ],
        stopping_sections=[
            "education",
            "skills",
            "technical skills",
            "projects",
            "certifications",
            "references",
        ],
    )

def extract_projects(text: str) -> str | None:
    return extract_section(
        text=text,
        section_names=[
            "projects",
            "academic projects",
            "personal projects",
            "project experience",
        ],
        stopping_sections=[
            "education",
            "experience",
            "work experience",
            "skills",
            "certifications",
            "references",
        ],
    )

def parse_cv_bytes(
    file_bytes: bytes,
    filename: str,
) -> dict[str, Any]:
    """
    Parse a CV and return structured information.
    """
    if not filename:
        raise CVParserError("A filename is required.")

    extension = Path(filename).suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        raise CVParserError(
            "Unsupported file type. Only PDF, DOCX and TXT files are allowed."
        )

    if not file_bytes:
        raise CVParserError("The uploaded file is empty.")

    if extension == ".pdf":
        text = extract_text_from_pdf(file_bytes)

    elif extension == ".docx":
        text = extract_text_from_docx(file_bytes)

    if not text:
        raise CVParserError(
            "No readable text was found in the uploaded CV."
        )

    return {
        "filename": filename,
        "file_type": extension.removeprefix("."),
        "character_count": len(text),
        "word_count": len(text.split()),
        "email": extract_email(text),
        "phone": extract_phone(text),
        "skills": extract_skills(text),
        "education": extract_education(text),
        "experience": extract_experience(text),
        "projects": extract_projects(text),
        "raw_text": text,
    }